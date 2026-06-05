"""P2-11: Export story to EPUB / DOCX / Markdown formats."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from worldforger.schemas import World
from worldforger.story.story_store import (
    macro_outline_path,
    manuscript_path,
    read_text,
    sorted_chapters,
)


def _strip_markdown_for_plain(text: str) -> str:
    """Remove common markdown formatting for plain text export."""
    # Remove headers markers but keep the text
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Remove bold/italic
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    # Remove code blocks
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _markdown_to_html(text: str) -> str:
    """Basic Markdown-to-HTML conversion for EPUB chapters."""
    # Headers
    text = re.sub(r"^#### (.+)$", r"<h4>\1</h4>", text, flags=re.MULTILINE)
    text = re.sub(r"^### (.+)$", r"<h3>\1</h3>", text, flags=re.MULTILINE)
    text = re.sub(r"^## (.+)$", r"<h2>\1</h2>", text, flags=re.MULTILINE)
    text = re.sub(r"^# (.+)$", r"<h1>\1</h1>", text, flags=re.MULTILINE)
    # Bold and italic
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
    # Horizontal rules
    text = re.sub(r"^---$", r"<hr/>", text, flags=re.MULTILINE)
    # Paragraphs: wrap lines separated by blank lines
    paragraphs = text.split("\n\n")
    result_parts = []
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if para.startswith("<h") or para.startswith("<hr"):
            result_parts.append(para)
        else:
            result_parts.append(f"<p>{para.replace(chr(10), '<br/>')}</p>")
    return "\n".join(result_parts)


def export_epub(world: World) -> bytes:
    """Generate an EPUB file for the world's story. Returns the binary content."""
    try:
        from ebooklib import epub
    except ImportError:
        raise ImportError(
            "ebooklib is required for EPUB export. Install with: pip install ebooklib"
        )

    wid = world.meta.id
    title = world.meta.name or "Untitled"
    chapters = sorted_chapters(world)
    macro = read_text(macro_outline_path(wid))

    book = epub.EpubBook()
    book.set_identifier(wid)
    book.set_title(title)
    book.set_language("zh-CN")
    book.add_author("Magic Creater World")

    # CSS
    style = """
    body { font-family: serif; line-height: 1.8; margin: 2em; }
    h1 { text-align: center; margin-bottom: 1em; }
    h2 { margin-top: 1.5em; }
    h3 { margin-top: 1em; }
    p { text-indent: 2em; margin: 0.5em 0; }
    """
    css = epub.EpubItem(
        uid="style", file_name="style/default.css", media_type="text/css", content=style
    )
    book.add_item(css)

    # Title page
    title_html = f"<h1>{title}</h1><p style='text-align:center'>by Magic Creater World</p>"
    title_chap = epub.EpubHtml(
        title="Title Page", file_name="title.xhtml", lang="zh-CN"
    )
    title_chap.content = f"<html><body>{title_html}</body></html>"
    book.add_item(title_chap)

    # Macro outline (as preface)
    epub_chapters: list = [title_chap]
    if macro.strip():
        macro_html = _markdown_to_html(macro)
        macro_chap = epub.EpubHtml(
            title="故事粗纲", file_name="macro_outline.xhtml", lang="zh-CN"
        )
        macro_chap.content = (
            f"<html><body><h2>故事粗纲</h2>{macro_html}</body></html>"
        )
        book.add_item(macro_chap)
        epub_chapters.append(macro_chap)

    # Chapter files
    for ch in chapters:
        ch_text = read_text(manuscript_path(wid, ch.id))
        if not ch_text.strip():
            continue
        ch_html = _markdown_to_html(ch_text)
        ch_title = ch.title or f"Chapter {ch.order}"
        chap = epub.EpubHtml(
            title=ch_title,
            file_name=f"chapter_{ch.order}.xhtml",
            lang="zh-CN",
        )
        chap.content = f"<html><body>{ch_html}</body></html>"
        book.add_item(chap)
        epub_chapters.append(chap)

    # Table of contents
    book.toc = epub_chapters

    # Navigation
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    # Spine
    book.spine = ["nav"] + epub_chapters

    buf = BytesIO()
    epub.write_epub(buf, book)
    return buf.getvalue()


def export_docx(world: World) -> bytes:
    """Generate a DOCX file for the world's story. Returns the binary content."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX export. Install with: pip install python-docx"
        )

    wid = world.meta.id
    title = world.meta.name or "Untitled"
    chapters = sorted_chapters(world)
    macro = read_text(macro_outline_path(wid))

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = 1  # center

    # Macro outline
    if macro.strip():
        doc.add_heading("故事粗纲", level=1)
        for line in macro.splitlines():
            line = line.strip()
            if not line:
                doc.add_paragraph("")
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=2)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=3)
            else:
                p = doc.add_paragraph(_strip_markdown_for_plain(line))
                style = p.style
                style.font.size = Pt(11)

    # Chapters
    for ch in chapters:
        ch_text = read_text(manuscript_path(wid, ch.id))
        if not ch_text.strip():
            continue
        ch_title = ch.title or f"Chapter {ch.order}"
        doc.add_heading(ch_title, level=1)
        for line in ch_text.splitlines():
            line = line.strip()
            if not line:
                doc.add_paragraph("")
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=2)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=3)
            else:
                p = doc.add_paragraph(_strip_markdown_for_plain(line))
                style = p.style
                style.font.size = Pt(11)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_markdown(world: World) -> str:
    """Generate a combined Markdown file for the world's story."""
    wid = world.meta.id
    title = world.meta.name or "Untitled"
    chapters = sorted_chapters(world)
    macro = read_text(macro_outline_path(wid))

    parts = [f"# {title}\n"]
    if macro.strip():
        parts.append("## 故事粗纲\n")
        parts.append(macro)
        parts.append("\n---\n")

    for ch in chapters:
        ch_text = read_text(manuscript_path(wid, ch.id))
        if not ch_text.strip():
            continue
        parts.append(f"## {ch.title or f'Chapter {ch.order}'}\n")
        parts.append(ch_text)
        parts.append("\n")

    return "\n".join(parts)
